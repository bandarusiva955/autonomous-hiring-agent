"""
Resume Parser - Extract text from PDF and DOCX files.
"""
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Try to import PDF and DOCX parsers
try:
    import PyPDF2
except ImportError:
    logger.warning("PyPDF2 not installed. PDF parsing will be limited.")
    PyPDF2 = None

try:
    import docx
except ImportError:
    logger.warning("python-docx not installed. DOCX parsing will be limited.")
    docx = None


def extract_text_from_file(filepath: str) -> str:
    """
    Extract text from a resume file (PDF or DOCX).
    
    Args:
        filepath: Path to the resume file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file format is not supported
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    # Get file extension
    ext = Path(filepath).suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: .pdf, .docx")


def extract_text_from_pdf(filepath: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        filepath: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    if PyPDF2 is None:
        raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
    
    text = []
    
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        
        result = '\n'.join(text)
        logger.info(f"Extracted {len(result)} characters from PDF: {filepath}")
        return result
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise


def extract_text_from_docx(filepath: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        filepath: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    if docx is None:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    try:
        doc = docx.Document(filepath)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        result = '\n'.join(paragraphs)
        logger.info(f"Extracted {len(result)} characters from DOCX: {filepath}")
        return result
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise


def validate_resume_file(filepath: str) -> tuple:
    """
    Validate a resume file.
    
    Args:
        filepath: Path to the resume file
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    if not os.path.exists(filepath):
        return False, "File not found"
    
    ext = Path(filepath).suffix.lower()
    if ext not in ['.pdf', '.docx']:
        return False, f"Unsupported file format: {ext}"
    
    # Check file size (max 10MB)
    file_size = os.path.getsize(filepath)
    max_size = 10 * 1024 * 1024  # 10MB
    
    if file_size > max_size:
        return False, f"File too large: {file_size / (1024*1024):.1f}MB (max 10MB)"
    
    return True, None


if __name__ == "__main__":
    # Test the parser
    print("Testing Resume Parser...")
    
    # Test with sample paths
    test_files = [
        "test_resume.pdf",
        "test_resume.docx"
    ]
    
    for test_file in test_files:
        print(f"\n{test_file}:")
        if os.path.exists(test_file):
            try:
                text = extract_text_from_file(test_file)
                print(f"  Extracted {len(text)} characters")
            except Exception as e:
                print(f"  Error: {e}")
        else:
            print("  File not found (expected in test)")

